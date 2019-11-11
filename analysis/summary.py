#! /usr/bin/env python
# -*- coding:utf-8 -*-
"""
@author  : MG
@Time    : 19-9-6 上午7:24
@File    : summary.py
@contact : mmmaaaggg@163.com
@desc    : 
"""

import datetime
import logging
import os

import docx
import pandas as pd
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from ibats_common.analysis.plot import clean_cache
from ibats_common.analysis.summary import df_2_table, stats_df_2_docx_table, dic_2_table
from ibats_common.backend.mess import get_report_folder_path
from ibats_utils.mess import datetime_2_str, date_2_str

logger = logging.getLogger(__name__)

STR_FORMAT_DATETIME_4_FILE_NAME = '%Y-%m-%d %H_%M_%S'
FORMAT_2_PERCENT = lambda x: f"{x * 100: .2f}%"
FORMAT_2_FLOAT2 = r"{0:.2f}"
FORMAT_2_FLOAT4 = r"{0:.4f}"
FORMAT_2_MONEY = r"{0:,.2f}"


def summary_analysis_result_dic_2_docx(round_results_dic: dict, title_header,
                                       enable_clean_cache=False, doc_file_path=None, ignore_if_exist=False):
    """对各个 round 分析结果进行汇总生产 docx 文件"""
    # 生成文件名
    if doc_file_path is not None:
        if os.path.isdir(doc_file_path):
            folder_path, file_name = doc_file_path, ''
        else:
            folder_path, file_name = os.path.split(doc_file_path)
    else:
        folder_path, file_name = None, ''

    if folder_path is None or folder_path == "":
        folder_path = get_report_folder_path()

    if folder_path != '' and not os.path.exists(folder_path):
        os.makedirs(folder_path)

    if file_name == '':
        file_name = f"{title_header}_" \
            f"{datetime_2_str(datetime.datetime.now(), STR_FORMAT_DATETIME_4_FILE_NAME)}.docx"
        file_path = os.path.join(folder_path, file_name)
    else:
        file_path = doc_file_path

    if ignore_if_exist and os.path.exists(file_path):
        return file_path

    # 生成 docx 文档将所需变量
    heading_title = f'Rewards 汇总分析报告[{title_header}] '

    # 生成 docx 文件
    document = docx.Document()
    # 设置默认字体
    document.styles['Normal'].font.name = '微软雅黑'
    document.styles['Normal']._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), '微软雅黑')
    # 创建自定义段落样式(第一个参数为样式名, 第二个参数为样式类型, 1为段落样式, 2为字符样式, 3为表格样式)
    UserStyle1 = document.styles.add_style('UserStyle1', 1)
    # 设置字体尺寸
    UserStyle1.font.size = docx.shared.Pt(40)
    # 设置字体颜色
    UserStyle1.font.color.rgb = docx.shared.RGBColor(0xff, 0xde, 0x00)
    # 居中文本
    UserStyle1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 设置中文字体
    UserStyle1.font.name = '微软雅黑'
    UserStyle1._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), '微软雅黑')

    # 文件内容
    heading_count, sub_heading_count = 0, 0
    document.add_heading(heading_title, 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph('')
    document.add_paragraph('')

    heading_count += 1
    sub_heading_count = 0
    key = 'available_episode_model_path_dic'
    document.add_heading(f'{heading_count}、各轮次训练有效的 episode 模型路径', 1)
    available_round_set = set()
    for num, (round_n, result_dic) in enumerate(round_results_dic.items()):
        analysis_result_dic = result_dic['analysis_result_dic']
        if key in analysis_result_dic:
            key_count = len(analysis_result_dic[key])
            sub_heading_count += 1
            run = document.add_heading('', 2).add_run(
                f'{heading_count}.{sub_heading_count}、第 {round_n} 轮训练[ {key_count} ]个')
            if key_count == 0:
                run.font.double_strike = True
            else:
                available_round_set.add(round_n)
            dic_2_table(document, analysis_result_dic[key], col_group_num=1, col_width=[0.6, 5.4])

    document.add_page_break()
    heading_count += 1
    document.add_heading(f'{heading_count}、各轮次训练 episode 与 value 变化趋势', 1)
    for num, (round_n, result_dic) in enumerate(round_results_dic.items()):
        run = document.add_heading('', 2).add_run(f'{heading_count}.{num}、Round {round_n} ')
        if round_n not in available_round_set:
            run.font.double_strike = True
        sub_heading_count = 0
        key = 'param_dic'
        if key in result_dic:
            # 将测试运行参数输出到 word
            param_dic = result_dic[key]
            if param_dic is not None and len(param_dic) > 0:
                sub_heading_count += 1
                document.add_heading(f'{heading_count}.{num}.{sub_heading_count}、相关参数 ', 3)
                document.add_paragraph("""\t包括训练、模型路径等相关初始化参数""")
                format_dic = {}
                dic_2_table(document, param_dic, col_group_num=1, format_dic=format_dic)

        key, key1 = 'analysis_result_dic', 'episode_in_sample_value_plot'
        if key in result_dic and key1 in result_dic[key]:
            sub_heading_count += 1
            document.add_heading(f'{heading_count}.{num}.{sub_heading_count}、趋势变化 ', 3)
            image_path = result_dic[key][key1]
            document.add_picture(image_path)  # , width=docx.shared.Inches(1.25)
            if num > 0:  # 第一段前面由于多了一个二级标题，导致最后一行分页时总是多出一个空白页
                document.add_page_break()

    # 保存文件
    document.save(file_path)
    if enable_clean_cache:
        clean_cache()

    return file_path


def in_out_example_analysis_result_all_round_2_docx(model_param_dic,
                                                    round_results_dic: dict, title_header, enable_clean_cache=False,
                                                    doc_file_path=None, ignore_if_exist=False):
    """对各个 round 分析结果进行汇总生产 docx 文件"""
    # 相关参数
    analysis_result_dic_key = 'analysis_result_dic'
    # 生成文件名
    if doc_file_path is not None:
        if os.path.isdir(doc_file_path):
            folder_path, file_name = doc_file_path, ''
        else:
            folder_path, file_name = os.path.split(doc_file_path)
    else:
        folder_path, file_name = None, ''

    if folder_path is None or folder_path == "":
        folder_path = get_report_folder_path()

    if folder_path != '' and not os.path.exists(folder_path):
        os.makedirs(folder_path)

    if file_name == '':
        file_name = f"{title_header}_" \
            f"{datetime_2_str(datetime.datetime.now(), STR_FORMAT_DATETIME_4_FILE_NAME)}.docx"
        docx_file_path = os.path.join(folder_path, file_name)
    else:
        docx_file_path = doc_file_path

    if ignore_if_exist and os.path.exists(docx_file_path):
        return docx_file_path

    # 生成 docx 文档将所需变量
    heading_title = f'Rewards 汇总分析报告[{title_header}] '

    # 生成 docx 文件
    document = docx.Document()
    # 设置默认字体
    document.styles['Normal'].font.name = '微软雅黑'
    document.styles['Normal']._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), '微软雅黑')
    # 创建自定义段落样式(第一个参数为样式名, 第二个参数为样式类型, 1为段落样式, 2为字符样式, 3为表格样式)
    UserStyle1 = document.styles.add_style('UserStyle1', 1)
    # 设置字体尺寸
    UserStyle1.font.size = docx.shared.Pt(40)
    # 设置字体颜色
    UserStyle1.font.color.rgb = docx.shared.RGBColor(0xff, 0xde, 0x00)
    # 居中文本
    UserStyle1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 设置中文字体
    UserStyle1.font.name = '微软雅黑'
    UserStyle1._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), '微软雅黑')

    # 文件内容
    title_num_list = [0, 0, 0]
    document.add_heading(heading_title, 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph('')
    document.add_paragraph('')

    # 将测试运行参数输出到 word
    if model_param_dic is not None and len(model_param_dic) > 0:
        title_num_list[0] += 1
        document.add_heading(f'{title_num_list[0]}、相关参数', 1)
        document.add_paragraph("""\t包括训练、模型路径等相关初始化参数""")
        format_dic = {}
        dic_2_table(document, model_param_dic, col_group_num=1, format_dic=format_dic)

    title_num_list[0] += 1
    title_num_list[1], title_num_list[2] = 0, 0
    document.add_heading(f'{title_num_list[0]}、各轮次训练有效的 episode 模型路径', 1)
    document.add_paragraph("""本节展示经过一定筛选逻辑后找到的有效模型及路径，其中样本外有效模型路径不具有实际操作价值，仅供参考""")
    key = 'available_episode_model_path_dic'
    available_round_set = set()
    for num, (round_n, result_dic) in enumerate(round_results_dic.items()):
        analysis_result_dic = result_dic[analysis_result_dic_key]
        for in_out_key in ('in_example', 'off_example'):
            in_out_key_str = '样本内' if in_out_key == 'in_example' else '样本外'
            if in_out_key not in analysis_result_dic:
                continue
            title_num_list[1] += 1
            data_dic = analysis_result_dic[in_out_key][key]
            data_count = len(data_dic)
            run = document.add_heading('', 2).add_run(
                f'{title_num_list[0]}、{title_num_list[1]} {in_out_key_str} 第 {round_n} 轮训练[ {data_count} ]个')
            if data_count == 0:
                run.font.double_strike = True
            else:
                available_round_set.add(round_n)

            dic_2_table(document, data_dic, col_group_num=1, col_width=[0.6, 5.4])
            document.add_paragraph('')

    document.add_page_break()

    title_num_list[0] += 1
    title_num_list[1], title_num_list[2] = 0, 0
    document.add_heading(f'{title_num_list[0]}、各轮次训练 episode 与 value 变化趋势', 1)
    for num, (round_n, result_dic) in enumerate(round_results_dic.items()):
        run = document.add_heading('', 2).add_run(f'{title_num_list[0]}.{num}、Round {round_n} ')
        if round_n not in available_round_set:
            run.font.double_strike = True
        title_num_list[1] = 0
        key, key1 = 'analysis_kwargs', 'model_param_dic'
        if key in result_dic and key1 in result_dic[key]:
            # 将测试运行参数输出到 word
            model_param_dic = result_dic[key][key1]
            if model_param_dic is not None and len(model_param_dic) > 0:
                title_num_list[1] += 1
                document.add_heading(f'{title_num_list[0]}.{num}.{title_num_list[1]}、Round {round_n} 相关参数 ', 3)
                document.add_paragraph("""\t包括训练、模型路径等相关初始化参数""")
                format_dic = {}
                dic_2_table(document, model_param_dic, col_group_num=1, format_dic=format_dic)

        key, key2 = 'analysis_result_dic', 'episode_nav_plot_file_path_dic'
        if key in result_dic:
            for in_out_key in ('in_example', 'off_example'):
                if in_out_key not in result_dic[key]:
                    continue
                in_out_key_str = '样本内' if in_out_key == 'in_example' else '样本外'
                title_num_list[1] += 1
                document.add_heading(
                    f'{title_num_list[0]}.{num}.{title_num_list[1]}、Round {round_n} {in_out_key_str} 趋势变化 ', 3)
                n_days_file_path_dic = result_dic[key][in_out_key][key2]
                for n_days, image_path in n_days_file_path_dic.items():
                    document.add_heading(
                        f'{title_num_list[0]}.{num}.{title_num_list[1]}、Round {round_n} {n_days} 日净值走势变化 ', 4)
                    document.add_picture(image_path)  # , width=docx.shared.Inches(1.25)

                if num > 0:  # 第一段前面由于多了一个二级标题，导致最后一行分页时总是多出一个空白页
                    document.add_page_break()

                # 保存文件
                document.save(docx_file_path)
                if enable_clean_cache:
                    clean_cache()

    return docx_file_path


def summary_rewards_2_docx(model_param_dic, analysis_result_dic, title_header, enable_clean_cache=False,
                           doc_file_path=None, in_sample_date_line=None):
    """创建 绩效分析结果 word 文档"""
    # 整理参数
    in_sample_date_line = None if in_sample_date_line is None else pd.to_datetime(in_sample_date_line)
    in_sample_date_line_str = date_2_str(in_sample_date_line) if in_sample_date_line is not None else ""
    int_col_name_set = {'action', 'action_count'}
    money_col_name_set = {'value', 'cash', 'fee_tot', 'value_fee0'}
    float_col_name_set = {'value', 'cash', 'close', 'fee_tot', 'value_fee0'}
    available_reward_col_name_list = ['value', 'cash', 'fee_tot', 'value_fee0', 'action_count']

    # 生成 docx 文档将所需变量
    heading_title = f'Rewards 分析报告[{title_header}] '

    # 生成 docx 文件
    document = docx.Document()
    # 设置默认字体
    document.styles['Normal'].font.name = '微软雅黑'
    document.styles['Normal']._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), '微软雅黑')
    # 创建自定义段落样式(第一个参数为样式名, 第二个参数为样式类型, 1为段落样式, 2为字符样式, 3为表格样式)
    UserStyle1 = document.styles.add_style('UserStyle1', 1)
    # 设置字体尺寸
    UserStyle1.font.size = docx.shared.Pt(40)
    # 设置字体颜色
    UserStyle1.font.color.rgb = docx.shared.RGBColor(0xff, 0xde, 0x00)
    # 居中文本
    UserStyle1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 设置中文字体
    UserStyle1.font.name = '微软雅黑'
    UserStyle1._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), '微软雅黑')

    # 文件内容
    heading_count, sub_heading_count = 0, 0
    document.add_heading(heading_title, 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph('')
    document.add_paragraph('')

    # 将测试运行参数输出到 word
    if model_param_dic is not None and len(model_param_dic) > 0:
        heading_count += 1
        document.add_heading(f'{heading_count}、相关参数', 1)
        document.add_paragraph("""\t包括训练、模型路径等相关初始化参数""")
        format_dic = {}
        dic_2_table(document, model_param_dic, col_group_num=1, format_dic=format_dic)

    # 随 Episode 增长，样本内 value 与 样本外 5日、10日、20日、60日 value 趋势对比
    # 目的是查看 样本内表现与样本外表现直接是否存在一定的相关性
    heading_count += 1
    document.add_heading(f'{heading_count}、样本内与样本外 5日、10日、20日、60日 Episode-value 趋势对比', 1)
    document.add_paragraph("""本节展示随 Episode 增长，样本内 value 与 样本外 5日、10日、20日、60日 value 趋势对比，
    目的是查看 样本内表现与样本外表现直接是否存在一定的相关性""")
    key = 'episode_value_compare'
    if key in analysis_result_dic:
        for n_days, file_path in analysis_result_dic[key].items():
            sub_heading_count += 1
            document.add_heading(f'{heading_count}、{sub_heading_count} 样本内与样本外{n_days}日对比走势', 1)
            document.add_picture(file_path)  # , width=docx.shared.Inches(1.25)
            document.add_paragraph('')

    # 样本内与样本外 value 相关性

    # 展示训练曲线 随 Episode 增长，value 结果变化曲线
    heading_count += 1
    document.add_heading(f'{heading_count}、展示训练曲线', 1)
    key = 'episode_in_sample_value_df'
    if key in analysis_result_dic:
        document.add_heading(f'{heading_count}.{sub_heading_count}、随 Episode 增长，value 结果变化曲线（样本内）'
                             f'{in_sample_date_line_str}', 2)
        document.add_picture(analysis_result_dic['episode_in_sample_value_plot'])
        document.add_paragraph('')

    key = 'episode_value_df'
    if key in analysis_result_dic:
        sub_heading_count += 1
        document.add_heading(f'{heading_count}.{sub_heading_count}、随 Episode 增长，value 结果变化曲线', 2)
        document.add_picture(analysis_result_dic['episode_value_plot'])  # , width=docx.shared.Inches(1.25)
        document.add_paragraph('详细数据')
        data_df = analysis_result_dic[key]

        format_by_col = {_: FORMAT_2_FLOAT4 for _ in data_df.columns if _ in float_col_name_set}
        format_by_col.update({_: FORMAT_2_PERCENT for _ in data_df.columns if _.find('rr') > 0 or _.find('cagr') > 0})
        df_2_table(document, data_df, format_by_col=format_by_col, max_col_count=6,
                   mark_top_n=5, mark_top_n_on_cols=data_df.columns)
        heading_count += 1
        # 添加分页符
        document.add_page_break()

    # 展示 Value Plot
    document.add_heading(f'{heading_count}、Reward Value 曲线走势', 1)
    # 汇总展示
    sub_heading_count = 1
    document.add_heading(f'{heading_count}.{sub_heading_count}、合并绘图走势', 2)
    for num, file_path in enumerate(analysis_result_dic['value_plot_list'], start=1):
        document.add_paragraph(f'{heading_count}.{sub_heading_count}.{num}')
        document.add_picture(file_path)
        document.add_paragraph('')
    data_df = analysis_result_dic['episode_reward_df'][available_reward_col_name_list]
    format_by_col = {_: FORMAT_2_FLOAT4 for _ in data_df.columns if _ in float_col_name_set}
    df_2_table(document, data_df, format_by_col=format_by_col, max_col_count=6,
               mark_top_n=5, mark_top_n_on_cols=data_df.columns)
    document.add_paragraph('')
    key = 'value_plot_dic'
    if key in analysis_result_dic:
        sub_heading_count += 1
        document.add_heading(f'{heading_count}.{sub_heading_count}、详细数据', 2)
        # 分别展示
        for num, (episode, file_path) in enumerate(analysis_result_dic[key].items(), start=2):
            document.add_heading(f'{heading_count}.{sub_heading_count}.{num}、episode={episode}', 2)
            document.add_picture(file_path)
            document.add_paragraph('')
            data_df = analysis_result_dic['episode_reward_dic'][episode][available_reward_col_name_list]
            format_by_col = {_: FORMAT_2_FLOAT4 for _ in data_df.columns if _ in float_col_name_set}
            df_2_table(document, data_df, format_by_col=format_by_col, max_col_count=6)
            document.add_paragraph('')

    heading_count += 1
    document.add_page_break()

    document.add_heading(f'{heading_count}、样本外数据绩效统计数据（Porformance stat）', 1)
    key = 'stats_df'
    if key in analysis_result_dic:
        stats_df = analysis_result_dic[key].T
        stats_df.drop(['start', 'rf'], axis=1, inplace=True)
        stats_df_2_docx_table(stats_df, document, format_axis='column',
                              mark_top_n=5, mark_top_n_on_cols=[_ for _ in stats_df.columns if _ not in {'end'}])
        heading_count += 1
        document.add_page_break()

    # 保存文件
    if doc_file_path is not None:
        if os.path.isdir(doc_file_path):
            folder_path, file_name = doc_file_path, ''
        else:
            folder_path, file_name = os.path.split(doc_file_path)
    else:
        folder_path, file_name = None, ''

    if folder_path is None or folder_path == "":
        folder_path = get_report_folder_path()

    if folder_path != '' and not os.path.exists(folder_path):
        os.makedirs(folder_path)

    if file_name == '':
        file_name = f"{title_header}_" \
            f"{datetime_2_str(datetime.datetime.now(), STR_FORMAT_DATETIME_4_FILE_NAME)}.docx"
        file_path = os.path.join(folder_path, file_name)
    else:
        file_path = doc_file_path

    document.save(file_path)
    if enable_clean_cache:
        clean_cache()

    return file_path


def in_out_example_analysis_result_2_docx(model_param_dic, analysis_result_dic, title_header, enable_clean_cache=False,
                                          doc_file_path=None, in_sample_date_line=None):
    """创建 绩效分析结果 word 文档"""
    # 整理参数
    in_sample_date_line = None if in_sample_date_line is None else pd.to_datetime(in_sample_date_line)
    in_sample_date_line_str = date_2_str(in_sample_date_line) if in_sample_date_line is not None else ""
    int_col_name_set = {'action', 'action_count'}
    money_col_name_set = {'value', 'cash', 'fee_tot', 'value_fee0'}
    float_col_name_set = {'nav', 'cash', 'close', 'fee_tot', 'nav_fee0'}
    available_reward_col_name_list = ['nav', 'nav_fee0', 'action_count', 'cash', 'fee_tot']

    # 生成 docx 文档将所需变量
    heading_title = f'分析报告[{title_header}] '

    # 生成 docx 文件
    document = docx.Document()
    # 设置默认字体
    document.styles['Normal'].font.name = '微软雅黑'
    document.styles['Normal']._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), '微软雅黑')
    # 创建自定义段落样式(第一个参数为样式名, 第二个参数为样式类型, 1为段落样式, 2为字符样式, 3为表格样式)
    UserStyle1 = document.styles.add_style('UserStyle1', 1)
    # 设置字体尺寸
    UserStyle1.font.size = docx.shared.Pt(40)
    # 设置字体颜色
    UserStyle1.font.color.rgb = docx.shared.RGBColor(0xff, 0xde, 0x00)
    # 居中文本
    UserStyle1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 设置中文字体
    UserStyle1.font.name = '微软雅黑'
    UserStyle1._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), '微软雅黑')

    # 文件内容
    title_num_list = [0, 0, 0]
    document.add_heading(heading_title, 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph('')
    document.add_paragraph('')

    # 将测试运行参数输出到 word
    if model_param_dic is not None and len(model_param_dic) > 0:
        title_num_list[0] += 1
        document.add_heading(f'{title_num_list[0]}、相关参数', 1)
        document.add_paragraph("""\t包括训练、模型路径等相关初始化参数""")
        format_dic = {}
        dic_2_table(document, model_param_dic, col_group_num=1, format_dic=format_dic)

    title_num_list[0] += 1
    title_num_list[1], title_num_list[2] = 0, 0
    document.add_heading(f'{title_num_list[0]}、有效的 episode 模型路径', 1)
    document.add_paragraph("""本节展示经过一定筛选逻辑后找到的有效模型及路径，其中样本外有效模型路径不具有实际操作价值，仅供参考""")
    key = 'available_episode_model_path_dic'
    for in_out_key in ('in_example', 'off_example'):
        in_out_key_str = '样本内' if in_out_key == 'in_example' else '样本外'
        if in_out_key not in analysis_result_dic or key not in analysis_result_dic[in_out_key]:
            continue
        title_num_list[1] += 1
        document.add_heading(f'{title_num_list[0]}、{title_num_list[1]} {in_out_key_str} 有效模型', 2)
        dic_2_table(document, analysis_result_dic[in_out_key][key], col_group_num=1, col_width=[0.6, 5.4])
        document.add_paragraph('')

    # 随 Episode 增长，nav 结果变化曲线
    title_num_list[0] += 1
    title_num_list[1], title_num_list[2] = 0, 0
    document.add_heading(f'{title_num_list[0]}、样本内与样本外 Episode-value 趋势对比', 1)
    document.add_paragraph("""本节展示随 Episode 增长，样本内 value 与 样本外 value 趋势对比，
    目的是查看 样本内表现与样本外表现直接是否存在一定的相关性""")
    key = 'episode_nav_plot_file_path_dic'
    for in_out_key in ('in_example', 'off_example'):
        in_out_key_str = '样本内' if in_out_key == 'in_example' else '样本外'
        if in_out_key not in analysis_result_dic or key not in analysis_result_dic[in_out_key]:
            continue
        title_num_list[1] += 1
        title_num_list[2] = 0
        document.add_heading(f'{title_num_list[0]}、{title_num_list[1]} {in_out_key_str} Episode NAV 走势变化', 2)
        n_days_file_path_dic = analysis_result_dic[in_out_key][key]
        for n_days, file_path in n_days_file_path_dic.items():
            title_num_list[2] += 1
            if n_days == -1:
                document.add_heading(
                    f'{title_num_list[0]}、{title_num_list[1]}、{title_num_list[2]} {in_out_key_str}最终净值', 3)
            else:
                document.add_heading(
                    f'{title_num_list[0]}、{title_num_list[1]}、{title_num_list[2]} {in_out_key_str}第{n_days}日净值', 3)
            document.add_picture(file_path)  # , width=docx.shared.Inches(1.25)
            document.add_paragraph('')

    # 随 Episode 增长，nav 结果变化曲线
    # 目的是查看 样本内表现与样本外表现直接是否存在一定的相关性
    title_num_list[0] += 1
    title_num_list[1], title_num_list[2] = 0, 0
    document.add_heading(f'{title_num_list[0]}、样本内与样本外 5日、10日、20日、60日 rr, cagr 变化列表', 1)
    document.add_paragraph("""本节展示随 Episode 增长，样本内 value 与 样本外 value 趋势对比，
    目的是查看 样本内表现与样本外表现直接是否存在一定的相关性""")
    key = 'episode_nav_df'
    for in_out_key in ('in_example', 'off_example'):
        in_out_key_str = '样本内' if in_out_key == 'in_example' else '样本外'
        if in_out_key not in analysis_result_dic or key not in analysis_result_dic[in_out_key]:
            continue
        title_num_list[1] += 1
        document.add_heading(f'{title_num_list[0]}、{title_num_list[1]} {in_out_key_str} 回测数据 rr,cagr 列表', 2)
        data_df = analysis_result_dic[in_out_key][key]
        format_by_col = {_: FORMAT_2_FLOAT4 for _ in data_df.columns if _ in float_col_name_set}
        format_by_col.update({_: FORMAT_2_PERCENT for _ in data_df.columns if _.find('rr') > 0 or _.find('cagr') > 0})
        df_2_table(document, data_df, format_by_col=format_by_col, max_col_count=6,
                   mark_top_n=5, mark_top_n_on_cols=data_df.columns)
        document.add_paragraph('')

    # 样本内与样本外 value 相关性
    key = 'episode_nav_plot_path_dic'
    title_num_list[0] += 1
    title_num_list[1], title_num_list[2] = 0, 0
    document.add_heading(f'{title_num_list[0]}、每一个 episode 分别展示 nav， nav_fee0 走势', 1)
    for in_out_key in ('in_example', 'off_example'):
        in_out_key_str = '样本内' if in_out_key == 'in_example' else '样本外'
        if in_out_key not in analysis_result_dic or key not in analysis_result_dic[in_out_key]:
            continue
        title_num_list[1] += 1
        title_num_list[2] = 0
        episode_nav_plot_path_dic = analysis_result_dic[in_out_key][key]
        data_count = len(episode_nav_plot_path_dic)
        document.add_heading(f'{title_num_list[0]}、{title_num_list[1]} {in_out_key_str} {data_count} 个 Episode 走势', 2)
        for episode, file_path in episode_nav_plot_path_dic.items():
            title_num_list[2] += 1
            document.add_heading(f'{title_num_list[0]}、{title_num_list[1]}、{title_num_list[2]} Episode {episode} 走势', 3)
            document.add_picture(file_path)  # , width=docx.shared.Inches(1.25)
            document.add_paragraph('')

    # 多 Episode 合并 nav 曲线走势
    key = 'nav_plot_file_path_list'
    title_num_list[0] += 1
    title_num_list[1], title_num_list[2] = 0, 0
    document.add_heading(f'{title_num_list[0]}、多 Episode 合并 nav 曲线走势', 1)
    for in_out_key in ('in_example', 'off_example'):
        in_out_key_str = '样本内' if in_out_key == 'in_example' else '样本外'
        if in_out_key not in analysis_result_dic or key not in analysis_result_dic[in_out_key]:
            continue
        title_num_list[1] += 1
        nav_plot_file_path_list = analysis_result_dic[in_out_key][key]
        data_count = len(nav_plot_file_path_list)
        document.add_heading(f'{title_num_list[0]}、{title_num_list[1]} {in_out_key_str} {data_count} 张合并后走势图', 2)
        for num, file_path in enumerate(nav_plot_file_path_list, start=1):
            document.add_paragraph(f'{title_num_list[0]}.{title_num_list[1]}.{num} {in_out_key_str} ')
            document.add_picture(file_path)
            document.add_paragraph('')

    key = 'episode_reward_df'
    title_num_list[0] += 1
    title_num_list[1], title_num_list[2] = 0, 0
    document.add_heading(f'{title_num_list[0]}、各 Episode Reward 列表', 1)
    for in_out_key in ('in_example', 'off_example'):
        in_out_key_str = '样本内' if in_out_key == 'in_example' else '样本外'
        if in_out_key not in analysis_result_dic or key not in analysis_result_dic[in_out_key]:
            continue
        title_num_list[1] += 1
        data_df = analysis_result_dic[in_out_key][key][available_reward_col_name_list]
        document.add_heading(f'{title_num_list[0]}、{title_num_list[1]} {in_out_key_str} 各Episode 最终净值列表', 2)
        format_by_col = {_: FORMAT_2_FLOAT4 for _ in data_df.columns if _ in float_col_name_set}
        df_2_table(document, data_df, format_by_col=format_by_col, max_col_count=6,
                   mark_top_n=5, mark_top_n_on_cols=data_df.columns)
        document.add_paragraph('')

    document.add_page_break()

    title_num_list[0] += 1
    key = 'stats_df'
    title_num_list[0] += 1
    title_num_list[1], title_num_list[2] = 0, 0
    document.add_heading(f'{title_num_list[0]}、样本外数据绩效统计数据（Porformance stat）', 1)
    for in_out_key in ('in_example', 'off_example'):
        in_out_key_str = '样本内' if in_out_key == 'in_example' else '样本外'
        if in_out_key not in analysis_result_dic or key not in analysis_result_dic[in_out_key]:
            continue
        title_num_list[1] += 1
        stats_df = analysis_result_dic[in_out_key][key].T
        document.add_heading(f'{title_num_list[0]}、{title_num_list[1]} {in_out_key_str} 绩效分析', 2)
        stats_df.drop(['start', 'rf'], axis=1, inplace=True)
        stats_df_2_docx_table(stats_df, document, format_axis='column',
                              mark_top_n=5, mark_top_n_on_cols=[_ for _ in stats_df.columns if _ not in {'end'}])
        title_num_list[0] += 1
        document.add_page_break()

    # 保存文件
    if doc_file_path is not None:
        if os.path.isdir(doc_file_path):
            folder_path, file_name = doc_file_path, ''
        else:
            folder_path, file_name = os.path.split(doc_file_path)
    else:
        folder_path, file_name = None, ''

    if folder_path is None or folder_path == "":
        folder_path = get_report_folder_path()

    if folder_path != '' and not os.path.exists(folder_path):
        os.makedirs(folder_path)

    if file_name == '':
        file_name = f"{title_header}_" \
            f"{datetime_2_str(datetime.datetime.now(), STR_FORMAT_DATETIME_4_FILE_NAME)}.docx"
        file_path = os.path.join(folder_path, file_name)
    else:
        file_path = doc_file_path

    document.save(file_path)
    if enable_clean_cache:
        clean_cache()

    return file_path
